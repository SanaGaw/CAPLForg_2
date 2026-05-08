"""STO (Specification Technique Object) parser for CAPL Pipeline V2.2.

Parses STO specification documents (.docx) to extract signal definitions
from tabular data including IN, OUT, Cal, Proxi, and DTC tables.
"""

from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, field
import logging
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


@dataclass
class StoSignal:
    """Represents a signal extracted from STO specification."""
    name: str
    sto_table_type: str  # 'IN', 'OUT', 'Cal', 'Proxi', 'DTC'
    ecu: Optional[str] = None
    value: Optional[str] = None
    unit: Optional[str] = None
    description: Optional[str] = None
    table_row: int = 0
    confidence: float = 0.8  # Base confidence for table extraction


@dataclass
class StoTable:
    """Represents a table extracted from STO specification."""
    table_type: str
    title: Optional[str] = None
    headers: List[str] = field(default_factory=list)
    signals: List[StoSignal] = field(default_factory=list)
    page_number: Optional[int] = None


class StoSpecParser:
    """
    Parse STO specification documents (.docx).

    Extracts ONLY tabular data from STO specifications.
    Ignores free-text content and schematic descriptions.

    Supported table types:
    - IN: Input signals
    - OUT: Output signals
    - Cal: Calibration parameters
    - Proxi: Proximity/sensor signals
    - DTC: Diagnostic trouble codes
    """

    # Table type indicators (case-insensitive)
    TABLE_TYPE_KEYWORDS = {
        'IN': ['input', 'in ', 'in-'],
        'OUT': ['output', 'out ', 'out-'],
        'Cal': ['calibration', 'cal ', 'kalibrier'],
        'Proxi': ['proxi', 'proximity', 'sensor'],
        'DTC': ['dtc', 'fault', 'diagnostic'],
    }

    # Keywords that indicate a table should be ignored
    IGNORE_KEYWORDS = [
        'figure', 'schema', 'abbildung', 'drawing',
        'note:', 'remark:', 'hinweis:',
    ]

    def __init__(self, docx_path: Path) -> None:
        self.docx_path = docx_path
        self._document: Optional[Document] = None
        self.tables: List[StoTable] = []
        self.signals: List[StoSignal] = []

    def _load_document(self) -> Document:
        """Load and cache the document."""
        if self._document is None:
            self._document = Document(self.docx_path)
        return self._document

    def parse(self) -> List[StoSignal]:
        """
        Parse STO document and extract signals from tables.

        Returns:
            List of StoSignal objects extracted from tables
        """
        self._load_document()
        self.tables = []
        self.signals = []

        for table in self._document.tables:
            sto_table = self._parse_table(table)
            if sto_table and sto_table.signals:
                self.tables.append(sto_table)
                self.signals.extend(sto_table.signals)

        logger.info(
            f"Parsed {self.docx_path.name}: "
            f"{len(self.tables)} tables, "
            f"{len(self.signals)} signals"
        )
        return self.signals

    def _parse_table(self, table: Table) -> Optional[StoTable]:
        """Parse a single table and determine its type."""
        rows = table.rows
        if len(rows) < 2:
            return None  # Need at least header + 1 data row

        # Get first row as potential header
        header_row = [cell.text.strip() for cell in rows[0].cells]
        header_text = ' '.join(header_row).lower()

        # Check if table should be ignored (not a signal table)
        if any(kw in header_text for kw in self.IGNORE_KEYWORDS):
            return None

        # Determine table type
        table_type = self._determine_table_type(header_text, header_row)

        if table_type is None:
            return None  # Not a recognized STO table type

        # Parse signals from table
        signals = self._extract_signals_from_table(table, table_type, header_row)

        return StoTable(
            table_type=table_type,
            headers=header_row,
            signals=signals
        )

    def _determine_table_type(self, header_text: str, header_row: List[str]) -> Optional[str]:
        """Determine the STO table type from header content."""
        header_text_lower = header_text.lower()

        for table_type, keywords in self.TABLE_TYPE_KEYWORDS.items():
            for keyword in keywords:
                if keyword in header_text_lower:
                    return table_type

        # Also check first column for table type markers
        if header_row and header_row[0]:
            first_cell = header_row[0].upper()
            for table_type in self.TABLE_TYPE_KEYWORDS:
                if table_type in first_cell:
                    return table_type

        return None

    def _extract_signals_from_table(
        self,
        table: Table,
        table_type: str,
        headers: List[str]
    ) -> List[StoSignal]:
        """Extract signal entries from table data rows."""
        signals = []
        rows = table.rows

        # Find name column (usually first or column with "name" in header)
        name_col = 0
        for i, header in enumerate(headers):
            if 'name' in header.lower() or 'signal' in header.lower():
                name_col = i
                break

        for row_idx, row in enumerate(rows[1:], start=2):
            cells = row.cells
            if not cells or len(cells) <= name_col:
                continue

            name_cell = cells[name_col].text.strip()
            if not name_cell or name_cell.lower() in ('-', 'n/a', 'none', ''):
                continue  # Skip empty rows

            # Create signal entry
            signal = StoSignal(
                name=name_cell,
                sto_table_type=table_type,
                table_row=row_idx
            )

            # Extract additional info from other columns
            if len(cells) > 1:
                signal.ecu = cells[1].text.strip() if cells[1].text.strip() else None

            if len(cells) > 2:
                val_cell = cells[2].text.strip()
                signal.value = val_cell if val_cell and val_cell not in ('-', 'n/a') else None

            if len(cells) > 3:
                signal.unit = cells[3].text.strip() if cells[3].text.strip() else None

            if len(cells) > 4:
                signal.description = cells[4].text.strip() if cells[4].text.strip() else None

            signals.append(signal)

        return signals

    def get_signals_by_type(self, table_type: str) -> List[StoSignal]:
        """Get all signals of a specific type."""
        if not self.signals:
            self.parse()
        return [s for s in self.signals if s.sto_table_type == table_type]

    def get_all_table_types(self) -> List[str]:
        """Get all table types found in document."""
        if not self.tables:
            self.parse()
        return list(set(t.table_type for t in self.tables))

    def generate_extract_report(self) -> dict:
        """Generate a structured report of the extraction."""
        if not self.signals:
            self.parse()

        report = {
            'source_file': str(self.docx_path),
            'total_tables': len(self.tables),
            'total_signals': len(self.signals),
            'by_type': {},
            'inconsistency_rate': 0.0,  # Will be calculated by cross-validator
        }

        for table_type in self.TABLE_TYPE_KEYWORDS:
            type_signals = self.get_signals_by_type(table_type)
            if type_signals:
                report['by_type'][table_type] = len(type_signals)

        return report
